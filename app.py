from flask import Flask, render_template, request, jsonify, send_file, session
from flask_cors import CORS
import os
import uuid
from werkzeug.utils import secure_filename
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import fitz  # PyMuPDF
import io
import base64
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
ALLOWED_EXTENSIONS = {'pdf'}

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)
os.makedirs('static/temp', exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 80 * 1024 * 1024  # 80MB max file size

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/favicon.ico')
def favicon():
    return '', 204  # No content

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        # Store in session
        session['current_pdf'] = unique_filename
        
        return jsonify({
            'success': True,
            'filename': unique_filename,
            'message': 'File uploaded successfully'
        })
    
    return jsonify({'error': 'Invalid file type. Please upload a PDF file.'}), 400

@app.route('/preview/<filename>')
def preview_pdf(filename):
    try:
        # Check both upload and processed folders (processed first for edited files)
        filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            app.logger.error(f'File not found: {filename}')
            return jsonify({'error': f'File not found: {filename}'}), 404
        
        app.logger.info(f'Loading preview for: {filename} from {filepath}')
        
        # Convert PDF pages to images for preview
        doc = fitz.open(filepath)
        pages = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))  # Higher resolution
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode()
            
            pages.append({
                'page_num': page_num + 1,
                'image': f"data:image/png;base64,{img_base64}",
                'width': pix.width,
                'height': pix.height
            })
        
        doc.close()
        
        response = jsonify({
            'success': True,
            'pages': pages,
            'total_pages': len(pages),
            'filename': filename,
            'timestamp': str(uuid.uuid4())  # Unique identifier for cache busting
        })
        
        # Add cache control headers to prevent caching of preview data
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        
        return response
        
    except Exception as e:
        return jsonify({'error': f'Error processing PDF: {str(e)}'}), 500

@app.route('/extract_text/<filename>')
def extract_text(filename):
    try:
        # Check both upload and processed folders
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            return jsonify({'error': f'File not found: {filename}'}), 404
        
        doc = fitz.open(filepath)
        pages_text = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            pages_text.append({
                'page_num': page_num + 1,
                'text': text
            })
        
        doc.close()
        
        return jsonify({
            'success': True,
            'pages_text': pages_text
        })
        
    except Exception as e:
        app.logger.error(f'Error extracting text from {filename}: {str(e)}')
        return jsonify({'error': f'Error extracting text: {str(e)}'}), 500

@app.route('/get_text_blocks/<filename>')
def get_text_blocks(filename):
    try:
        # Check both upload and processed folders
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            app.logger.error(f'File not found: {filename}')
            return jsonify({'error': f'File not found: {filename}'}), 404
        
        doc = fitz.open(filepath)
        pages_blocks = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_blocks = []
            
            try:
                # Try to get structured text blocks first
                blocks = page.get_text("dict")
                
                for block in blocks["blocks"]:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                if span["text"].strip():  # Only include non-empty text
                                    text_blocks.append({
                                        'text': span["text"],
                                        'bbox': span["bbox"],  # [x0, y0, x1, y1]
                                        'font': span.get("font", "unknown"),
                                        'size': span.get("size", 12),
                                        'flags': span.get("flags", 0),
                                        'color': span.get("color", 0)
                                    })
                
                # If no structured text found, try alternative methods
                if not text_blocks:
                    # Try getting text with bboxes using a different method
                    text_dict = page.get_text("rawdict")
                    for block in text_dict.get("blocks", []):
                        if block.get("type") == 0:  # Text block
                            for line in block.get("lines", []):
                                for span in line.get("spans", []):
                                    if span.get("text", "").strip():
                                        text_blocks.append({
                                            'text': span["text"],
                                            'bbox': span["bbox"],
                                            'font': span.get("font", "unknown"),
                                            'size': span.get("size", 12),
                                            'flags': span.get("flags", 0),
                                            'color': span.get("color", 0)
                                        })
                
                # If still no text, try simple text extraction with manual bbox estimation
                if not text_blocks:
                    simple_text = page.get_text()
                    if simple_text.strip():
                        # Create a single text block for the entire page content
                        page_rect = page.rect
                        text_blocks.append({
                            'text': simple_text,
                            'bbox': [50, 50, page_rect.width - 50, page_rect.height - 50],
                            'font': "unknown",
                            'size': 12,
                            'flags': 0,
                            'color': 0
                        })
                
            except Exception as text_error:
                # Fallback for problematic pages
                try:
                    simple_text = page.get_text()
                    if simple_text.strip():
                        page_rect = page.rect
                        text_blocks.append({
                            'text': simple_text,
                            'bbox': [50, 50, page_rect.width - 50, page_rect.height - 50],
                            'font': "unknown",
                            'size': 12,
                            'flags': 0,
                            'color': 0
                        })
                except:
                    pass  # Skip problematic pages
            
            pages_blocks.append({
                'page_num': page_num + 1,
                'blocks': text_blocks
            })
        
        doc.close()
        
        return jsonify({
            'success': True,
            'pages_blocks': pages_blocks
        })
        
    except Exception as e:
        app.logger.error(f'Error getting text blocks from {filename}: {str(e)}')
        return jsonify({'error': f'Error getting text blocks: {str(e)}'}), 500

@app.route('/add_text', methods=['POST'])
def add_text():
    try:
        data = request.json
        filename = data.get('filename')
        page_num = data.get('page_num', 1)
        text = data.get('text', '')
        x = data.get('x', 100)
        y = data.get('y', 100)
        font_size = data.get('font_size', 12)
        color = data.get('color', [0, 0, 0])  # RGB color
        
        if not filename:
            return jsonify({'error': 'No filename provided'}), 400
        
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        # Create new PDF with added text
        doc = fitz.open(filepath)
        page = doc.load_page(page_num - 1)
        
        # Get page dimensions for coordinate conversion
        page_rect = page.rect
        page_height = page_rect.height
        
        # Convert coordinates - PDF coordinate system has origin at bottom-left
        # but user interface expects top-left origin
        pdf_x = x
        pdf_y = page_height - y  # Flip Y coordinate
        
        # Create text rectangle for better positioning
        text_rect = fitz.Rect(pdf_x, pdf_y - font_size, pdf_x + 200, pdf_y + 5)
        
        # Insert text with better formatting
        page.insert_textbox(
            text_rect,
            text,
            fontsize=font_size,
            color=color,
            fontname="helv",  # Helvetica font
            align=0  # Left align
        )
        
        # Save modified PDF
        output_filename = f"modified_{uuid.uuid4()}_{filename}"
        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        doc.save(output_path)
        doc.close()
        
        return jsonify({
            'success': True,
            'modified_filename': output_filename,
            'message': 'Text added successfully',
            'coordinates': {'x': pdf_x, 'y': pdf_y, 'original_y': y}
        })
        
    except Exception as e:
        return jsonify({'error': f'Error adding text: {str(e)}'}), 500

@app.route('/edit_text', methods=['POST'])
def edit_text():
    try:
        data = request.json
        app.logger.info(f'Edit text request: {data}')
        
        filename = data.get('filename')
        page_num = data.get('page_num', 1)
        old_text = data.get('old_text', '')
        new_text = data.get('new_text', '')
        bbox = data.get('bbox')  # [x0, y0, x1, y1]
        
        if not filename:
            app.logger.error('No filename provided')
            return jsonify({'error': 'No filename provided'}), 400
            
        if not bbox or len(bbox) != 4:
            app.logger.error(f'Invalid bbox: {bbox}')
            return jsonify({'error': 'Invalid bounding box coordinates'}), 400
        
        # Check both upload and processed folders
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            app.logger.error(f'File not found: {filename}')
            return jsonify({'error': f'File not found: {filename}'}), 404
        
        app.logger.info(f'Opening PDF: {filepath}')
        
        # Open the PDF
        doc = fitz.open(filepath)
        
        if page_num < 1 or page_num > len(doc):
            doc.close()
            app.logger.error(f'Invalid page number: {page_num}')
            return jsonify({'error': f'Invalid page number: {page_num}'}), 400
        
        page = doc.load_page(page_num - 1)
        
        # Validate and adjust bbox coordinates if needed
        page_rect = page.rect
        bbox[0] = max(0, min(bbox[0], page_rect.width))
        bbox[1] = max(0, min(bbox[1], page_rect.height))
        bbox[2] = max(bbox[0], min(bbox[2], page_rect.width))
        bbox[3] = max(bbox[1], min(bbox[3], page_rect.height))
        
        # Create a white rectangle to cover the old text with minimal padding
        padding = 1  # Minimal padding to avoid covering adjacent text
        rect = fitz.Rect(
            max(0, bbox[0] - padding), 
            max(0, bbox[1] - padding), 
            min(page_rect.width, bbox[2] + padding), 
            min(page_rect.height, bbox[3] + padding)
        )
        page.draw_rect(rect, color=None, fill=(1, 1, 1))  # White fill
        app.logger.info(f'Covered old text area: {rect}')
        
        # Add the new text at the same position if provided
        if new_text and new_text.strip():
            # Get font preservation settings
            font_info = data.get('font_info', {})
            preserve_formatting = data.get('preserve_formatting', True)
            
            if preserve_formatting and font_info:
                # Use original font properties
                font_size = font_info.get('size', 12)
                font_name = font_info.get('font', 'Arial')
                font_color = font_info.get('color', 0)
                font_flags = font_info.get('flags', 0)
                
                app.logger.info(f'Using preserved formatting - Font: {font_name}, Size: {font_size}, Color: {font_color}, Flags: {font_flags}')
                
                # Convert font color from integer to RGB tuple
                if isinstance(font_color, int):
                    if font_color == 0:
                        color_rgb = (0, 0, 0)  # Black
                    else:
                        # Extract RGB components from integer (BGR format in PyMuPDF)
                        b = (font_color & 0xFF) / 255.0
                        g = ((font_color >> 8) & 0xFF) / 255.0
                        r = ((font_color >> 16) & 0xFF) / 255.0
                        color_rgb = (r, g, b)
                else:
                    color_rgb = (0, 0, 0)  # Default to black
                
                # Enhanced font mapping with better support for common fonts
                base_font = 'helv'  # Default fallback
                font_name_lower = font_name.lower()
                
                # Map common font families
                if any(name in font_name_lower for name in ['arial', 'helvetica']):
                    base_font = 'helv'
                elif any(name in font_name_lower for name in ['times', 'roman']):
                    base_font = 'times'
                elif any(name in font_name_lower for name in ['courier', 'mono']):
                    base_font = 'cour'
                elif 'symbol' in font_name_lower:
                    base_font = 'symb'
                elif 'zapf' in font_name_lower:
                    base_font = 'zadb'
                
                # Apply style flags more robustly
                is_bold = bool(font_flags & (1 << 4)) or 'bold' in font_name_lower
                is_italic = bool(font_flags & (1 << 6)) or any(style in font_name_lower for style in ['italic', 'oblique'])
                
                if is_bold and is_italic:
                    if base_font == 'times':
                        base_font = 'times-bolditalic'
                    elif base_font == 'helv':
                        base_font = 'helv-boldoblique'
                    elif base_font == 'cour':
                        base_font = 'cour-boldoblique'
                elif is_bold:
                    base_font += '-bold'
                elif is_italic:
                    if base_font == 'times':
                        base_font += '-italic'
                    else:
                        base_font += '-oblique'
                
                app.logger.info(f'Mapped font: {font_name} -> {base_font}')
            else:
                # Use default formatting
                bbox_height = bbox[3] - bbox[1]
                font_size = min(12, max(8, bbox_height * 0.7))
                base_font = 'helv'
                color_rgb = (0, 0, 0)
            
            # Calculate text position - use more accurate positioning
            text_x = bbox[0]
            text_y = bbox[3] - 2  # Slightly above the bottom of the original text bbox
            
            # Try multiple insertion methods with better error handling
            text_inserted = False
            
            # Method 1: Try with preserved font
            if preserve_formatting and font_info:
                try:
                    page.insert_text(
                        (text_x, text_y), 
                        new_text, 
                        fontsize=font_size, 
                        color=color_rgb,
                        fontname=base_font
                    )
                    app.logger.info(f'Text inserted with preserved font: {base_font}')
                    text_inserted = True
                except Exception as font_error:
                    app.logger.warning(f'Preserved font insertion failed: {font_error}')
            
            # Method 2: Try with basic helvetica
            if not text_inserted:
                try:
                    page.insert_text(
                        (text_x, text_y), 
                        new_text, 
                        fontsize=font_size, 
                        color=color_rgb,
                        fontname='helv'
                    )
                    app.logger.info('Text inserted with helvetica fallback')
                    text_inserted = True
                except Exception as helv_error:
                    app.logger.warning(f'Helvetica insertion failed: {helv_error}')
            
            # Method 3: Try with no font specification
            if not text_inserted:
                try:
                    page.insert_text(
                        (text_x, text_y), 
                        new_text, 
                        fontsize=font_size, 
                        color=(0, 0, 0)
                    )
                    app.logger.info('Text inserted with default font')
                    text_inserted = True
                except Exception as default_error:
                    app.logger.warning(f'Default font insertion failed: {default_error}')
            
            # Method 4: Last resort - use textbox
            if not text_inserted:
                try:
                    text_rect = fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3])
                    page.insert_textbox(
                        text_rect,
                        new_text,
                        fontsize=font_size,
                        color=(0, 0, 0),
                        align=0
                    )
                    app.logger.info('Text inserted using textbox fallback')
                    text_inserted = True
                except Exception as textbox_error:
                    app.logger.error(f'All text insertion methods failed: {textbox_error}')
            
            if not text_inserted:
                raise Exception("Failed to insert text using any method")
        
        # Generate a shorter, more manageable filename
        base_name = os.path.basename(filename)
        
        # If filename is too long, truncate it but keep the extension
        if len(base_name) > 50:
            name_part, ext = os.path.splitext(base_name)
            base_name = name_part[:40] + "..." + ext
        
        # Create output filename with timestamp for uniqueness
        import time
        timestamp = int(time.time())
        output_filename = f"edited_{timestamp}_{base_name}"
        
        # Ensure the full path isn't too long for Windows (260 char limit)
        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        if len(output_path) > 250:  # Leave some buffer
            # Use just timestamp and extension if still too long
            name_part, ext = os.path.splitext(base_name)
            output_filename = f"edited_{timestamp}{ext}"
            output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        
        app.logger.info(f'Saving edited PDF to: {output_path}')
        
        # Ensure the output directory exists
        os.makedirs(app.config['PROCESSED_FOLDER'], exist_ok=True)
        
        # Try to save with different methods if the first fails
        save_successful = False
        
        try:
            # Method 1: Standard save
            doc.save(output_path)
            save_successful = True
            app.logger.info('PDF saved successfully with standard method')
        except Exception as save_error:
            app.logger.warning(f'Standard save failed: {save_error}')
            
            # Method 2: Try with different save options
            try:
                doc.save(output_path, garbage=0, clean=False, deflate=False)
                save_successful = True
                app.logger.info('PDF saved successfully with alternative options')
            except Exception as alt_save_error:
                app.logger.warning(f'Alternative save failed: {alt_save_error}')
                
                # Method 3: Try saving to a temporary file first, then move
                try:
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                        temp_path = temp_file.name
                    
                    doc.save(temp_path)
                    
                    # Move the temp file to the final location
                    import shutil
                    shutil.move(temp_path, output_path)
                    save_successful = True
                    app.logger.info('PDF saved successfully using temporary file method')
                    
                except Exception as temp_save_error:
                    app.logger.error(f'All save methods failed: {temp_save_error}')
                    # Clean up temp file if it exists
                    try:
                        if 'temp_path' in locals() and os.path.exists(temp_path):
                            os.unlink(temp_path)
                    except:
                        pass
        
        doc.close()
        
        if not save_successful:
            raise Exception("Failed to save PDF using any method")
        
        return jsonify({
            'success': True,
            'modified_filename': output_filename,
            'message': 'Text edited successfully'
        })
        
    except Exception as e:
        app.logger.error(f'Error editing text: {str(e)}', exc_info=True)
        return jsonify({'error': f'Error editing text: {str(e)}'}), 500

@app.route('/merge_pdfs', methods=['POST'])
def merge_pdfs():
    try:
        data = request.json
        filenames = data.get('filenames', [])
        
        if len(filenames) < 2:
            return jsonify({'error': 'At least 2 files required for merging'}), 400
        
        merger = PyPDF2.PdfMerger()
        
        for filename in filenames:
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(filepath):
                merger.append(filepath)
        
        output_filename = f"merged_{uuid.uuid4()}.pdf"
        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        
        with open(output_path, 'wb') as output_file:
            merger.write(output_file)
        
        merger.close()
        
        return jsonify({
            'success': True,
            'merged_filename': output_filename,
            'message': 'PDFs merged successfully'
        })
        
    except Exception as e:
        return jsonify({'error': f'Error merging PDFs: {str(e)}'}), 500

@app.route('/split_pdf', methods=['POST'])
def split_pdf():
    try:
        data = request.json
        filename = data.get('filename')
        start_page = data.get('start_page', 1)
        end_page = data.get('end_page')
        
        if not filename:
            return jsonify({'error': 'No filename provided'}), 400
        
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        with open(filepath, 'rb') as input_file:
            reader = PyPDF2.PdfReader(input_file)
            writer = PyPDF2.PdfWriter()
            
            total_pages = len(reader.pages)
            if not end_page:
                end_page = total_pages
            
            for page_num in range(start_page - 1, min(end_page, total_pages)):
                writer.add_page(reader.pages[page_num])
            
            output_filename = f"split_{start_page}-{end_page}_{uuid.uuid4()}_{filename}"
            output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
        
        return jsonify({
            'success': True,
            'split_filename': output_filename,
            'message': f'PDF split successfully (pages {start_page}-{end_page})'
        })
        
    except Exception as e:
        return jsonify({'error': f'Error splitting PDF: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        # Check both upload and processed folders
        filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
            
    except Exception as e:
        return jsonify({'error': f'Error downloading file: {str(e)}'}), 500

@app.route('/download_word/<filename>')
def download_word_file(filename):
    """Download Word document files"""
    try:
        # Word files are always in processed folder
        filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        
        if os.path.exists(filepath):
            return send_file(
                filepath, 
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return jsonify({'error': 'Word file not found'}), 404
            
    except Exception as e:
        return jsonify({'error': f'Error downloading Word file: {str(e)}'}), 500

@app.route('/ocr_text/<filename>')
def ocr_text(filename):
    """Simple OCR text extraction for image-based PDFs"""
    try:
        # Check both upload and processed folders
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            app.logger.error(f'File not found for OCR: {filename}')
            return jsonify({'error': f'File not found: {filename}'}), 404
        
        doc = fitz.open(filepath)
        pages_text = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # First try regular text extraction
            text = page.get_text()
            
            # If no text found, provide a message about OCR
            if not text.strip():
                text = "This appears to be an image-based PDF. OCR functionality requires additional setup (Tesseract OCR). For now, you can still add new text to this document using the 'Add Text' feature."
            
            pages_text.append({
                'page_num': page_num + 1,
                'text': text,
                'method': 'basic_extraction'
            })
        
        doc.close()
        
        return jsonify({
            'success': True,
            'pages_text': pages_text,
            'ocr_used': False,
            'message': 'Basic text extraction completed. For full OCR support, install Tesseract OCR.'
        })
        
    except Exception as e:
        app.logger.error(f'Error in text extraction from {filename}: {str(e)}')
        return jsonify({'error': f'Error in text extraction: {str(e)}'}), 500

@app.route('/convert_to_word/<filename>')
def convert_to_word(filename):
    """Convert PDF to Word document with formatting preservation"""
    try:
        # Check both upload and processed folders
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            app.logger.error(f'File not found for Word conversion: {filename}')
            return jsonify({'error': f'File not found: {filename}'}), 404
        
        app.logger.info(f'Converting PDF to Word: {filename}')
        
        # Open PDF document
        doc = fitz.open(filepath)
        
        # Create Word document
        word_doc = Document()
        
        # Set document margins (narrower for better content fit)
        sections = word_doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Process each page
        total_text_blocks = 0
        for page_num in range(len(doc)):
            try:
                page = doc.load_page(page_num)
                
                # Add page break for pages after the first
                if page_num > 0:
                    word_doc.add_page_break()
                
                # Add page header (optional, can be disabled for cleaner output)
                if len(doc) > 1:  # Only add page numbers for multi-page documents
                    page_header = word_doc.add_paragraph()
                    page_header.add_run(f"Page {page_num + 1}").bold = True
                    page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    page_header.space_after = Pt(6)
                
                # Extract text blocks with formatting
                text_blocks = _extract_formatted_text_blocks(page)
                total_text_blocks += len(text_blocks)
                
                if text_blocks:
                    # Group text blocks by approximate lines
                    lines = _group_text_blocks_into_lines(text_blocks)
                    
                    # Convert lines to Word paragraphs
                    for line in lines:
                        _add_line_to_word_doc(word_doc, line)
                else:
                    # Fallback: extract simple text
                    simple_text = page.get_text()
                    if simple_text.strip():
                        # Split into paragraphs more intelligently
                        paragraphs = []
                        current_para = []
                        
                        for line in simple_text.split('\n'):
                            line = line.strip()
                            if line:
                                current_para.append(line)
                            else:
                                if current_para:
                                    paragraphs.append(' '.join(current_para))
                                    current_para = []
                        
                        # Add the last paragraph
                        if current_para:
                            paragraphs.append(' '.join(current_para))
                        
                        # Add paragraphs to Word document
                        for para_text in paragraphs:
                            if para_text.strip():
                                para = word_doc.add_paragraph(para_text.strip())
                                para.space_after = Pt(6)
                
                # Add some spacing between pages (but not after the last page)
                if page_num < len(doc) - 1:
                    word_doc.add_paragraph()
                    
            except Exception as page_error:
                app.logger.error(f'Error processing page {page_num + 1}: {str(page_error)}')
                # Add error message to document
                error_para = word_doc.add_paragraph()
                error_run = error_para.add_run(f"[Error processing page {page_num + 1}: {str(page_error)}]")
                error_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                error_run.italic = True
        
        # Store page count before closing document
        total_pages = len(doc)
        doc.close()
        
        # Generate output filename
        base_name = os.path.splitext(os.path.basename(filename))[0]
        if len(base_name) > 50:
            base_name = base_name[:50] + "..."
        
        import time
        timestamp = int(time.time())
        output_filename = f"converted_{timestamp}_{base_name}.docx"
        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        
        # Save Word document
        word_doc.save(output_path)
        
        app.logger.info(f'PDF converted to Word successfully: {output_filename}')
        
        return jsonify({
            'success': True,
            'word_filename': output_filename,
            'message': 'PDF converted to Word document successfully',
            'pages_processed': total_pages,
            'text_blocks_processed': total_text_blocks,
            'output_path': output_path,
            'file_size': os.path.getsize(output_path) if os.path.exists(output_path) else 0
        })
        
    except Exception as e:
        app.logger.error(f'Error converting PDF to Word: {str(e)}', exc_info=True)
        return jsonify({'error': f'Error converting PDF to Word: {str(e)}'}), 500

def _extract_formatted_text_blocks(page):
    """Extract text blocks with detailed formatting information"""
    text_blocks = []
    
    try:
        # Get structured text data
        text_dict = page.get_text("dict")
        
        for block in text_dict.get("blocks", []):
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        if span["text"].strip():
                            # Extract formatting details
                            font_info = {
                                'text': span["text"],
                                'bbox': span["bbox"],
                                'font': span.get("font", "Arial"),
                                'size': span.get("size", 12),
                                'flags': span.get("flags", 0),
                                'color': span.get("color", 0),
                                'origin': span.get("origin", [0, 0])
                            }
                            
                            # Determine text properties
                            font_info['is_bold'] = bool(font_info['flags'] & (1 << 4))
                            font_info['is_italic'] = bool(font_info['flags'] & (1 << 6))
                            font_info['is_underline'] = bool(font_info['flags'] & (1 << 2))
                            
                            # Convert color to RGB
                            color_int = font_info['color']
                            if color_int == 0:
                                font_info['rgb_color'] = (0, 0, 0)  # Black
                            else:
                                b = (color_int & 0xFF)
                                g = ((color_int >> 8) & 0xFF)
                                r = ((color_int >> 16) & 0xFF)
                                font_info['rgb_color'] = (r, g, b)
                            
                            text_blocks.append(font_info)
    
    except Exception as e:
        app.logger.warning(f'Error extracting formatted text blocks: {str(e)}')
    
    return text_blocks

def _group_text_blocks_into_lines(text_blocks):
    """Group text blocks into logical lines based on Y coordinates"""
    if not text_blocks:
        return []
    
    # Sort blocks by Y coordinate (top to bottom), then X coordinate (left to right)
    sorted_blocks = sorted(text_blocks, key=lambda b: (b['bbox'][1], b['bbox'][0]))
    
    lines = []
    current_line = []
    current_y = None
    y_tolerance = 5  # Pixels tolerance for same line
    
    for block in sorted_blocks:
        block_y = block['bbox'][1]  # Top Y coordinate
        
        if current_y is None or abs(block_y - current_y) <= y_tolerance:
            # Same line
            current_line.append(block)
            current_y = block_y
        else:
            # New line
            if current_line:
                lines.append(current_line)
            current_line = [block]
            current_y = block_y
    
    # Add the last line
    if current_line:
        lines.append(current_line)
    
    return lines

def _add_line_to_word_doc(word_doc, line_blocks):
    """Add a line of text blocks to Word document with formatting"""
    if not line_blocks:
        return
    
    # Create paragraph
    paragraph = word_doc.add_paragraph()
    
    # Determine paragraph alignment based on text position
    page_width = 595  # Approximate A4 width in points
    avg_x = sum(block['bbox'][0] for block in line_blocks) / len(line_blocks)
    
    if avg_x < page_width * 0.2:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif avg_x > page_width * 0.8:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif page_width * 0.4 < avg_x < page_width * 0.6:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add text runs with formatting
    for block in line_blocks:
        run = paragraph.add_run(block['text'])
        
        # Apply font formatting
        try:
            # Font size
            run.font.size = Pt(max(8, min(72, block['size'])))
            
            # Font family (map PDF fonts to Word fonts)
            font_name = _map_pdf_font_to_word(block['font'])
            run.font.name = font_name
            
            # Font style
            run.font.bold = block['is_bold']
            run.font.italic = block['is_italic']
            run.font.underline = block['is_underline']
            
            # Font color
            r, g, b = block['rgb_color']
            run.font.color.rgb = RGBColor(r, g, b)
            
        except Exception as e:
            app.logger.warning(f'Error applying formatting to text run: {str(e)}')
        
        # Add space between blocks if they're not adjacent
        if block != line_blocks[-1]:  # Not the last block
            next_block = line_blocks[line_blocks.index(block) + 1]
            current_right = block['bbox'][2]
            next_left = next_block['bbox'][0]
            
            # Add space if there's a gap
            if next_left - current_right > 5:  # 5 pixel threshold
                paragraph.add_run(' ')

def _map_pdf_font_to_word(pdf_font):
    """Map PDF font names to Word-compatible font names"""
    font_mapping = {
        # Helvetica family
        'helv': 'Arial',
        'helvetica': 'Arial',
        'helveticabold': 'Arial',
        'helvetica-bold': 'Arial',
        'helveticaneue': 'Arial',
        'helveticaneueb': 'Arial',
        
        # Times family
        'times': 'Times New Roman',
        'times-roman': 'Times New Roman',
        'times-bold': 'Times New Roman',
        'times-italic': 'Times New Roman',
        'times-bolditalic': 'Times New Roman',
        'timesnewroman': 'Times New Roman',
        
        # Courier family
        'cour': 'Courier New',
        'courier': 'Courier New',
        'courier-bold': 'Courier New',
        'courier-oblique': 'Courier New',
        'couriernew': 'Courier New',
        
        # Other common fonts
        'arial': 'Arial',
        'calibri': 'Calibri',
        'verdana': 'Verdana',
        'georgia': 'Georgia',
        'trebuchet': 'Trebuchet MS',
    }
    
    # Clean and normalize font name
    clean_font = pdf_font.lower().replace('-', '').replace(' ', '')
    
    # Try exact match first
    if clean_font in font_mapping:
        return font_mapping[clean_font]
    
    # Try partial matches
    for pdf_name, word_name in font_mapping.items():
        if pdf_name in clean_font or clean_font in pdf_name:
            return word_name
    
    # Default fallback
    return 'Arial'

def cleanup_old_files():
    """Clean up old processed files to prevent disk space issues"""
    try:
        import time
        current_time = time.time()
        max_age = 24 * 60 * 60  # 24 hours in seconds
        
        for folder in [app.config['PROCESSED_FOLDER'], app.config['UPLOAD_FOLDER']]:
            if os.path.exists(folder):
                for filename in os.listdir(folder):
                    filepath = os.path.join(folder, filename)
                    if os.path.isfile(filepath):
                        file_age = current_time - os.path.getmtime(filepath)
                        if file_age > max_age:
                            try:
                                os.unlink(filepath)
                                app.logger.info(f'Cleaned up old file: {filename}')
                            except Exception as e:
                                app.logger.warning(f'Failed to clean up {filename}: {e}')
    except Exception as e:
        app.logger.error(f'Cleanup process failed: {e}')

@app.route('/debug_pdf/<filename>')
def debug_pdf(filename):
    """Debug endpoint to check PDF content after editing"""
    try:
        # Check both folders
        filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
        if not os.path.exists(filepath):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        doc = fitz.open(filepath)
        debug_info = {
            'filename': filename,
            'filepath': filepath,
            'page_count': len(doc),
            'pages': []
        }
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            text_dict = page.get_text("dict")
            
            page_info = {
                'page_num': page_num + 1,
                'text_length': len(text),
                'text_preview': text[:200] + '...' if len(text) > 200 else text,
                'text_blocks_count': len([b for b in text_dict.get("blocks", []) if "lines" in b])
            }
            debug_info['pages'].append(page_info)
        
        doc.close()
        return jsonify(debug_info)
        
    except Exception as e:
        return jsonify({'error': f'Debug error: {str(e)}'}), 500

if __name__ == '__main__':
    # Clean up old files on startup
    cleanup_old_files()
    app.run(debug=True, host='0.0.0.0', port=5000)